from googletrans import Translator
import os
from docx import Document
from werkzeug.utils import secure_filename

translator = Translator()

UPLOAD_FOLDER = 'uploads'
TRANSLATED_FOLDER = 'translated_files'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(TRANSLATED_FOLDER, exist_ok=True)

ALLOWED_EXTENSIONS = {'docx'}

languages = {
    'kn': 'Kannada',
    'hi': 'Hindi',
    'ml': 'Malayalam',
    'ta': 'Tamil'
}


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def read_word_file(filepath):
    doc = Document(filepath)
    return '\n'.join([para.text for para in doc.paragraphs])


def write_word_file(filepath, text):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(filepath)


def translate_and_save(file, paragraph, language):
    translated_text = ''
    translated_filename = ''

    if file:
        filename = secure_filename(file.filename)
        filepath = os.path.join(UPLOAD_FOLDER, filename)
        file.save(filepath)

        if filename.endswith('.docx'):
            paragraph = read_word_file(filepath)

    if not paragraph:
        paragraph = paragraph

    translation = translator.translate(paragraph, src='en', dest=language)
    translated_text = translation.text

    translated_filename = f'translated_{language}.docx'
    translated_filepath = os.path.join(TRANSLATED_FOLDER, translated_filename)
    write_word_file(translated_filepath, translated_text)

    return translated_text, paragraph, translated_filename
